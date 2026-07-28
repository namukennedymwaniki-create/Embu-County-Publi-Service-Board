# ai_knowledge_base.py
import streamlit as st
import google.generativeai as genai
import PyPDF2
import uuid
import io
from datetime import datetime
from typing import List, Dict
import psycopg2

try:
    from pgvector.psycopg2 import register_vector
except ImportError:
    print("⚠️ pgvector not available")

# Word document support
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not installed. Word documents will not be supported.")

class AIKnowledgeBase:
    def __init__(self):
        # Get API key from secrets
        self.api_key = st.secrets.get("GEMINI_API_KEY")
        if not self.api_key:
            st.error("❌ GEMINI_API_KEY not found in secrets. Please add it.")
            return
        
        # Configure Gemini
        genai.configure(api_key=self.api_key)
        
        # =========================================================
        # CHAT MODELS - Use available models
        # =========================================================
        self.chat_models_to_try = [
            "models/gemini-2.5-flash",
            "models/gemini-2.0-flash",
            "models/gemini-flash-latest",
            "models/gemini-2.5-pro",
            "models/gemini-pro-latest",
        ]
        
        # Find working chat model
        self.chat_model = None
        for model_name in self.chat_models_to_try:
            try:
                test_model = genai.GenerativeModel(model_name)
                response = test_model.generate_content("Test")
                if response and response.text:
                    self.chat_model = model_name
                    print(f"✅ Using chat model: {model_name}")
                    break
            except Exception as e:
                print(f"❌ {model_name} failed: {e}")
                continue
        
        if not self.chat_model:
            st.error("❌ No working chat model found. Please check your API key.")
            return
        
        # =========================================================
        # EMBEDDING MODELS - Use available models
        # =========================================================
        self.embedding_models = [
            "models/gemini-embedding-2",
            "models/gemini-embedding-2-preview",
            "models/gemini-embedding-001",
        ]
        
        self.chunk_size = 1000
        self.chunk_overlap = 200
        
        print(f"✅ AI Knowledge Base initialized successfully")
        print(f"📊 Chat model: {self.chat_model}")
        print(f"📄 Word documents support: {'✅' if DOCX_AVAILABLE else '❌'}")
    
    def get_conn(self):
        """Get database connection with pgvector support"""
        database_url = st.secrets.get("DATABASE_URL")
        if database_url:
            try:
                conn = psycopg2.connect(database_url, sslmode='require')
                try:
                    register_vector(conn)
                except:
                    pass
                return conn
            except Exception as e:
                st.error(f"Database connection error: {e}")
                return None
        return None
    
    def extract_text_simple(self, file_content: bytes, filename: str) -> str:
        """Simple text extraction from multiple file formats"""
        filename_lower = filename.lower()
        text = ""
        
        try:
            # PDF files
            if filename_lower.endswith('.pdf'):
                import PyPDF2
                import io
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # Word documents
            elif filename_lower.endswith('.docx'):
                try:
                    import docx
                    import io
                    doc = docx.Document(io.BytesIO(file_content))
                    for para in doc.paragraphs:
                        if para.text.strip():
                            text += para.text + "\n"
                    # Extract from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    text += cell.text + " "
                            text += "\n"
                except ImportError:
                    st.error("❌ python-docx not installed. Run: pip install python-docx")
                    return ""
            
            # Text files
            elif filename_lower.endswith('.txt'):
                try:
                    text = file_content.decode('utf-8')
                except:
                    text = file_content.decode('latin-1')
            
            else:
                st.error(f"❌ Unsupported file format: {filename}")
                return ""
            
            # Clean up
            text = text.strip()
            
            if not text:
                st.warning("⚠️ No text extracted from document")
            else:
                st.success(f"✅ Extracted {len(text)} characters")
            
            return text
            
        except Exception as e:
            st.error(f"❌ Error extracting text: {e}")
            return ""
    
    def chunk_text(self, text: str) -> List[Dict]:
        """Split text into overlapping chunks"""
        chunks = []
        words = text.split()
        total_words = len(words)
        
        if total_words == 0:
            return chunks
        
        for i in range(0, total_words, self.chunk_size - self.chunk_overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = " ".join(chunk_words)
            chunks.append({
                'chunk_number': len(chunks) + 1,
                'chunk_text': chunk_text
            })
            
            if i + self.chunk_size >= total_words:
                break
        
        return chunks
    
    def create_embedding(self, text: str) -> List[float]:
        """Create embedding using Gemini with fallback models"""
        for model_name in self.embedding_models:
            try:
                result = genai.embed_content(
                    model=model_name,
                    content=text[:8191],
                    task_type="retrieval_document"
                )
                print(f"✅ Embedding successful with {model_name}")
                return result['embedding']
            except Exception as e:
                print(f"❌ {model_name} failed: {e}")
                continue
        
        st.error("❌ All embedding models failed")
        return []
    
    def process_document(self, file_content: bytes, filename: str, title: str, 
                         category: str, uploaded_by: str) -> Dict:
        """Process and store a document - supports PDF, Word, and Text files"""
        try:
            # Extract text from the document
            text = self.extract_text_simple(file_content, filename)
            
            if not text or not text.strip():
                return {
                    'success': False, 
                    'error': 'No text could be extracted from this document. Please ensure it contains readable text.'
                }
            
            st.info(f"📊 Extracted {len(text)} characters of text")
            
            # Show preview
            preview = text[:300] + "..." if len(text) > 300 else text
            st.info(f"📝 Preview: {preview}")
            
            # Create document record
            doc_id = str(uuid.uuid4())
            conn = self.get_conn()
            if not conn:
                return {'success': False, 'error': 'Database connection failed'}
            
            cursor = conn.cursor()
            
            # Save document record
            st.info("💾 Saving document record...")
            cursor.execute("""
                INSERT INTO documents (
                    id, filename, title, category, uploaded_by, 
                    file_size, page_count, is_active
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING id
            """, (
                doc_id, filename, title, category, uploaded_by,
                len(file_content), text.count('\n') // 40 + 1, True
            ))
            
            # Chunk the text
            st.info("✂️ Splitting text into chunks...")
            chunks = self.chunk_text(text)
            st.info(f"📋 Created {len(chunks)} chunks")
            
            if not chunks:
                return {'success': False, 'error': 'No chunks created from text'}
            
            # Create embeddings for each chunk
            chunk_count = 0
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            for i, chunk in enumerate(chunks):
                status_text.text(f"Processing chunk {i+1}/{len(chunks)}...")
                
                if i == 0:
                    st.info(f"📝 First chunk: {chunk['chunk_text'][:150]}...")
                
                embedding = self.create_embedding(chunk['chunk_text'])
                if embedding:
                    chunk_id = str(uuid.uuid4())
                    cursor.execute("""
                        INSERT INTO document_chunks (
                            id, document_id, chunk_number, chunk_text, embedding
                        ) VALUES (%s, %s, %s, %s, %s)
                    """, (
                        chunk_id, doc_id, chunk['chunk_number'],
                        chunk['chunk_text'], embedding
                    ))
                    chunk_count += 1
                else:
                    st.warning(f"⚠️ Failed to create embedding for chunk {i+1}")
                
                progress_bar.progress((i + 1) / len(chunks))
            
            conn.commit()
            conn.close()
            status_text.empty()
            progress_bar.empty()
            
            if chunk_count == 0:
                return {
                    'success': False,
                    'error': f'No embeddings created. Failed to process all {len(chunks)} chunks'
                }
            
            return {
                'success': True,
                'document_id': doc_id,
                'chunks_created': chunk_count,
                'total_chunks': len(chunks),
                'message': f'Document processed with {chunk_count}/{len(chunks)} chunks'
            }
            
        except Exception as e:
            st.error(f"❌ Error: {str(e)}")
            return {'success': False, 'error': str(e)}
    
    def search_documents(self, question: str, limit: int = 5) -> List[Dict]:
        """Search for relevant document chunks"""
        try:
            question_embedding = self.create_embedding(question)
            if not question_embedding:
                return []
            
            conn = self.get_conn()
            if not conn:
                return []
            
            cursor = conn.cursor()
            
            cursor.execute("""
                SELECT 
                    dc.chunk_text,
                    dc.page_number,
                    d.title as document_title,
                    d.filename,
                    1 - (dc.embedding <=> %s::vector) as similarity
                FROM document_chunks dc
                JOIN documents d ON dc.document_id = d.id
                WHERE d.is_active = TRUE
                ORDER BY dc.embedding <=> %s::vector
                LIMIT %s
            """, (question_embedding, question_embedding, limit))
            
            results = cursor.fetchall()
            conn.close()
            
            if not results:
                return []
            
            return [{
                'chunk_text': r[0],
                'page_number': r[1],
                'document_title': r[2],
                'filename': r[3],
                'similarity': r[4]
            } for r in results]
            
        except Exception as e:
            st.error(f"Search error: {e}")
            return []
    
    def generate_answer(self, question: str, chunks: List[Dict]) -> Dict:
        """Generate answer using Gemini"""
        try:
            if not chunks:
                return {
                    'answer': "I couldn't find information related to your question in the uploaded knowledge base.",
                    'sources': [],
                    'follow_up': [],
                    'confidence': 0.0
                }
            
            context = "\n\n".join([
                f"From document '{chunk['document_title']}' (page {chunk['page_number']}):\n{chunk['chunk_text']}"
                for chunk in chunks[:3]
            ])
            
            prompt = f"""
            You are an AI assistant for Embu County Public Service Board. 
            Answer the following question based ONLY on the provided context.
            If the answer cannot be found in the context, say "I couldn't find information related to your question in the uploaded knowledge base."
            Do not use any external knowledge or make assumptions.
            
            Question: {question}
            
            Context:
            {context}
            
            Answer:"""
            
            model = genai.GenerativeModel(self.chat_model)
            response = model.generate_content(prompt)
            
            answer = response.text
            
            sources = [
                {
                    'title': chunk['document_title'],
                    'page': chunk['page_number'],
                    'filename': chunk['filename']
                }
                for chunk in chunks[:3]
            ]
            
            # Calculate confidence based on similarity
            confidence = 0.0
            if chunks:
                similarities = [chunk.get('similarity', 0) for chunk in chunks[:3]]
                avg_similarity = sum(similarities) / len(similarities) if similarities else 0
                if avg_similarity > 0.75:
                    confidence = 0.9
                elif avg_similarity > 0.6:
                    confidence = 0.7
                elif avg_similarity > 0.4:
                    confidence = 0.5
                else:
                    confidence = 0.3
            
            return {
                'answer': answer,
                'sources': sources,
                'follow_up': [],
                'confidence': confidence
            }
            
        except Exception as e:
            return {
                'answer': f"Error generating answer: {str(e)}",
                'sources': [],
                'follow_up': [],
                'confidence': 0.0
            }
